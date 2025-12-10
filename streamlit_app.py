
import streamlit as st
import pandas as pd
from docx import Document
import zipfile
import io
from datetime import datetime
import time

# ============== CONFIGURACIÓN DE PÁGINA ==============
st.set_page_config(
    page_title="ATENEA: Generador Estudios Previos y Minutas",
    page_icon="Generador de Documentos",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ============== ESTILOS CSS PERSONALIZADOS ==============
st.markdown("""
<style>
    /* Fuentes y colores principales */
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');
    
    .main-header {
        background: linear-gradient(135deg, #1e3a5f 0%, #2d5a87 100%);
        padding: 2rem;
        border-radius: 16px;
        margin-bottom: 2rem;
        box-shadow: 0 4px 20px rgba(0,0,0,0.1);
    }
    
    .main-title {
        color: white;
        font-size: 2.5rem;
        font-weight: 700;
        margin: 0;
        display: flex;
        align-items: center;
        gap: 12px;
    }
    
    .subtitle {
        color: rgba(255,255,255,0.8);
        font-size: 1.1rem;
        margin-top: 0.5rem;
    }
    
    .upload-card {
        background: white;
        border: 2px dashed #e0e0e0;
        border-radius: 12px;
        padding: 1.5rem;
        text-align: center;
        transition: all 0.3s ease;
    }
    
    .upload-card:hover {
        border-color: #2d5a87;
        background: #f8fafc;
    }
    
    .stat-card {
        background: linear-gradient(135deg, #f8fafc 0%, #e2e8f0 100%);
        border-radius: 12px;
        padding: 1.5rem;
        text-align: center;
        border-left: 4px solid #2d5a87;
    }
    
    .stat-number {
        font-size: 2rem;
        font-weight: 700;
        color: #1e3a5f;
    }
    
    .stat-label {
        color: #64748b;
        font-size: 0.9rem;
    }
    
    .success-box {
        background: linear-gradient(135deg, #ecfdf5 0%, #d1fae5 100%);
        border-left: 4px solid #10b981;
        padding: 1rem 1.5rem;
        border-radius: 8px;
        margin: 1rem 0;
    }
    
    .info-box {
        background: linear-gradient(135deg, #eff6ff 0%, #dbeafe 100%);
        border-left: 4px solid #3b82f6;
        padding: 1rem 1.5rem;
        border-radius: 8px;
    }
    
    .step-badge {
        background: #2d5a87;
        color: white;
        padding: 4px 12px;
        border-radius: 20px;
        font-size: 0.8rem;
        font-weight: 600;
    }
    
    /* Ocultar elementos de Streamlit */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    
    /* Estilo para el sidebar */
    .css-1d391kg {
        background: linear-gradient(180deg, #1e3a5f 0%, #2d5a87 100%);
    }
    
    .sidebar-info {
        background: rgba(255,255,255,0.1);
        border-radius: 8px;
        padding: 1rem;
        margin: 1rem 0;
    }
</style>
""", unsafe_allow_html=True)

# ============== FUNCIONES ==============
def replace_text_in_paragraph(paragraph, key, value):
    """Reemplaza placeholders en párrafos manteniendo formato"""
    placeholder = "{{" + key + "}}"
    value_str = str(value) if value is not None else ""
    full_text = "".join(run.text for run in paragraph.runs)
    if placeholder in full_text:
        for run in paragraph.runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, value_str)
                return

def generar_documentos(df, word_file, progress_bar, status_text):
    """Genera documentos Word a partir de Excel con barra de progreso"""
    try:
        zip_buffer = io.BytesIO()
        documentos_generados = 0
        errores = []
        total = len(df)
        
        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zipf:
            for idx, row in df.iterrows():
                try:
                    word_file.seek(0)
                    doc = Document(word_file)
                    
                    for paragraph in doc.paragraphs:
                        for key in row.index:
                            replace_text_in_paragraph(paragraph, key, row[key])
                    
                    for table in doc.tables:
                        for row_table in table.rows:
                            for cell in row_table.cells:
                                for paragraph in cell.paragraphs:
                                    for key in row.index:
                                        replace_text_in_paragraph(paragraph, key, row[key])
                    
                    doc_bytes = io.BytesIO()
                    doc.save(doc_bytes)
                    zipf.writestr(f"Documento_{idx + 1}.docx", doc_bytes.getvalue())
                    documentos_generados += 1
                    
                    # Actualizar progreso
                    progress = (idx + 1) / total
                    progress_bar.progress(progress)
                    status_text.text(f"📄 Procesando documento {idx + 1} de {total}...")
                    
                except Exception as e:
                    errores.append(f"Fila {idx + 1}: {str(e)}")
        
        status_text.text("✅ ¡Proceso completado!")
        return zip_buffer, documentos_generados, errores
        
    except Exception as e:
        raise Exception(f"Error procesando documentos: {str(e)}")

# ============== SIDEBAR ==============
with st.sidebar:
    st.markdown("## Generador Inteligente de Documentos")
    st.markdown("---")
    
    st.markdown("### 📖 Guía Rápida")
    
    with st.expander("1️⃣ Preparar Excel", expanded=False):
        st.markdown("""
        - Crea un archivo Excel (.xlsx)
        - Cada columna será una variable
        - Cada fila genera un documento
        """)
    
    with st.expander("2️⃣ Preparar Plantilla", expanded=False):
        st.markdown("""
        - Crea un documento Word (.docx)
        - Usa `{{nombre_columna}}` para las variables
        - Ejemplo: `{{nombre}}`, `{{fecha}}`
        """)
    
    with st.expander("3️⃣ Generar", expanded=False):
        st.markdown("""
        - Carga ambos archivos
        - Haz clic en "Generar Documentos"
        - Descarga el ZIP con los resultados
        """)
    
    st.markdown("---")
    
    st.markdown("### ℹ️ Información")
    st.markdown("""
    <div class="sidebar-info">
        <p><strong>Versión:</strong> 1.0</p>
        <p><strong>Grencia de Gestión Corporativa</p>
    
                
    </div>
    """, unsafe_allow_html=True)
    

# ============== CONTENIDO PRINCIPAL ==============

# Header
st.markdown("""
<div class="main-header">
    <h1 class="main-title" style="color: white;">Generador Inteligente de Documentos</h1>
    <p class="subtitle">Estudios Previos y Minutas | Geh1rencia Gestión Corporativa</p>
</div>
""", unsafe_allow_html=True)

# Pasos del proceso
st.markdown("### 📋 Proceso de Generación")

col1, col2, col3 = st.columns(3)

with col1:
    st.markdown("""
    <div class="stat-card">
        <span class="step-badge">Paso 1</span>
        <h4 style="margin: 1rem 0 0.5rem 0;">📊 Cargar Datos</h4>
        <p style="color: #64748b; font-size: 0.9rem;">Sube tu archivo Excel con la información</p>
    </div>
    """, unsafe_allow_html=True)

with col2:
    st.markdown("""
    <div class="stat-card">
        <span class="step-badge">Paso 2</span>
        <h4 style="margin: 1rem 0 0.5rem 0;">📝 Cargar Plantilla</h4>
        <p style="color: #64748b; font-size: 0.9rem;">Sube tu plantilla Word con placeholders</p>
    </div>
    """, unsafe_allow_html=True)

with col3:
    st.markdown("""
    <div class="stat-card">
        <span class="step-badge">Paso 3</span>
        <h4 style="margin: 1rem 0 0.5rem 0;">🚀 Generar</h4>
        <p style="color: #64748b; font-size: 0.9rem;">Genera y descarga tus documentos</p>
    </div>
    """, unsafe_allow_html=True)

st.markdown("<br>", unsafe_allow_html=True)

# Área de carga de archivos
st.markdown("### 📁 Cargar Archivos")

col1, col2 = st.columns(2)

with col1:
    st.markdown("##### 📊 Datos (Excel)")
    excel_file = st.file_uploader(
        "Arrastra o selecciona tu archivo Excel",
        type="xlsx",
        key="excel",
        help="Archivo Excel con los datos para generar documentos"
    )
    if excel_file:
        st.success(f"✅ {excel_file.name}")

with col2:
    st.markdown("##### 📝 Plantilla (Word)")
    word_file = st.file_uploader(
        "Arrastra o selecciona tu plantilla Word",
        type="docx",
        key="word",
        help="Plantilla Word con placeholders {{columna}}"
    )
    if word_file:
        st.success(f"✅ {word_file.name}")

# Procesamiento
if excel_file and word_file:
    st.markdown("---")
    
    # Vista previa de datos
    with st.expander("👀 Vista previa de datos", expanded=True):
        try:
            df_preview = pd.read_excel(excel_file)
            
            # Métricas
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("📄 Documentos a generar", len(df_preview))
            with col2:
                st.metric("📊 Columnas/Variables", len(df_preview.columns))
            with col3:
                st.metric("🔤 Placeholders detectados", len(df_preview.columns))
            
            st.markdown("<br>", unsafe_allow_html=True)
            st.dataframe(df_preview, use_container_width=True, height=200)
            
            # Mostrar columnas disponibles
            st.markdown("**Variables disponibles para la plantilla:**")
            placeholders = " | ".join([f"`{{{{{col}}}}}`" for col in df_preview.columns])
            st.code(placeholders, language=None)
            
        except Exception as e:
            st.error(f"Error al leer Excel: {e}")
    
    st.markdown("<br>", unsafe_allow_html=True)
    
    # Botón de generación
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        generate_btn = st.button(
            "🚀 Generar Documentos",
            use_container_width=True,
            type="primary"
        )
    
    if generate_btn:
        try:
            excel_file.seek(0)
            df = pd.read_excel(excel_file)
            
            if df.empty:
                st.error("❌ El archivo Excel está vacío")
            else:
                st.markdown("---")
                st.markdown("### ⚙️ Procesando...")
                
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                zip_buffer, generados, errores = generar_documentos(
                    df, word_file, progress_bar, status_text
                )
                
                time.sleep(0.5)  # Pequeña pausa para mejor UX
                
                if generados > 0:
                    st.markdown("---")
                    
                    # Resultados
                    st.markdown("""
                    <div class="success-box">
                        <h3 style="margin: 0; color: #065f46;">🎉 ¡Generación Exitosa!</h3>
                    </div>
                    """, unsafe_allow_html=True)
                    
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        st.metric("✅ Generados", generados)
                    with col2:
                        st.metric("⚠️ Errores", len(errores))
                    with col3:
                        st.metric("📊 Total procesados", len(df))
                    
                    if errores:
                        with st.expander(f"⚠️ Ver {len(errores)} errores"):
                            for error in errores:
                                st.warning(error)
                    
                    st.markdown("<br>", unsafe_allow_html=True)
                    
                    # Botón de descarga
                    col1, col2, col3 = st.columns([1, 2, 1])
                    with col2:
                        zip_buffer.seek(0)
                        st.download_button(
                            label=f"📥 Descargar ZIP ({generados} documentos)",
                            data=zip_buffer.getvalue(),
                            file_name=f"Documentos_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip",
                            mime="application/zip",
                            use_container_width=True,
                            type="primary"
                        )
                else:
                    st.error("❌ No se pudieron generar documentos")
                    with st.expander("Ver errores"):
                        for error in errores:
                            st.error(error)
                            
        except Exception as e:
            st.error(f"❌ Error: {str(e)}")
            st.info("💡 Verifica que el Excel tenga datos y la plantilla tenga placeholders {{columna}}")

else:
    # Estado inicial - sin archivos
    st.markdown("---")
    st.markdown("""
    <div class="info-box">
        <h4 style="margin: 0 0 0.5rem 0;">👋 ¡Bienvenido!</h4>
        <p style="margin: 0; color: #1e40af;">
            Carga un archivo Excel y una plantilla Word para comenzar a generar documentos automáticamente.
        </p>
    </div>
    """, unsafe_allow_html=True)
    
    # Ejemplo de uso
    with st.expander("📚 Ver ejemplo de uso"):
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("**Ejemplo de Excel:**")
            ejemplo_df = pd.DataFrame({
                "Descripción": ["objeto", "objeto"],
                "Dependencia": ["GEP", "GGC"],
                "fecha": ["2024-01-15", "2024-01-16"]
            })
            st.dataframe(ejemplo_df, use_container_width=True)
        
        with col2:
            st.markdown("**Ejemplo de Plantilla:**")
            st.code("""
Certificado de Trabajo

Certificamos que {{nombre}} 
desempeña el cargo de {{cargo}}
desde el {{fecha}}.
            """)